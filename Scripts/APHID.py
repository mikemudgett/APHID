import sys, os
from Bio import SeqIO
import numpy as np
import pandas as pd
import primer3
from io import StringIO
from Bio.Seq import Seq
import regex as re

from Bio.Blast.Applications import NcbiblastnCommandline

from docx import Document
from docx.shared import RGBColor

from aphid_utils import Gene, pull_guides, grow_primer_to_tm, choose_primer

args = sys.argv
if len(args) != 5:
	print('Input Error.')
	print('')
	print('Usage: python APHID.py {Transcript ID} {Insertion Site} {Cargo File} {Project Name}')
	print('Example: python APHID.py AT4G31820.1 -1 GFP.fa NPY1_GFP')
	print('')
	print('Argument details:')
	print('Transcript ID: This is a TAIR10 transcript ID. Typically you will use the transcript from the representative gene model, or the transcript with the longest coding sequence. In most cases, this will be your gene ID plus ".1".')
	print('Insertion Site: This is the amino acid residue after which you would like your cargo insertion. For example, to insert after the initial M, use 1. To automatically insert at the end of the coding sequence, use -1.')
	print('Cargo File: A FASTA file with the nulceotide sequence of the cargo you would like to insert. This file should be in your local directory and should only have one entry, otherwise the last sequence will be used.')
	print('Project Name: A title for your project. This will be used to name output files.')
	sys.exit()



try:
	#Start with the name of the transcript we are interested in. Some genes have more than one so we need to specify. If it is not specified, we can add in a method to assume the user wants the '.1' version
	transcript_name = str(args[1])
	#Add something to check that it is a valid number.
	
	insertion_location = int(args[2])
	cargo_fn = str(args[3])
	cargo_name = cargo_fn.split('.')[0]
	project_name = str(args[4])


except:
	print('Input Error.')
	print('')
	print('Usage: python APHID.py {Transcript ID} {Insertion Site} {Cargo File} {Project Name}')
	print('Example: python APHID.py AT4G31820.1 -1 GFP.fa NPY1_GFP')
	print('')
	print('Argument details:')
	print('Transcript ID: This is a TAIR10 transcript ID. Typically you will use the transcript from the representative gene model, or the transcript with the longest coding sequence. In most cases, this will be your gene ID plus ".1".')
	print('Insertion Site: This is the amino acid residue after which you would like your cargo insertion. For example, to insert after the initial M, use 1. To automatically insert at the end of the coding sequence, use -1.')
	print('Cargo File: A FASTA file with the nulceotide sequence of the cargo you would like to insert. This file should be in your local directory and should only have one entry, otherwise the last sequence will be used.')
	print('Project Name: A title for your project. This will be used to name output files.')
	sys.exit()

#Specify this should have one sequence only.
#do this in a try as well:
try:
	cargo_file = open(cargo_fn, 'r')
except:
	print('Error! Could not find cargo file.')
	sys.exit()

#VARIABLES

local_seq_range = 1500 #How much sequence to look at on either side of the insertion site. Do not change unless you are using very large homology arms or want more sequence in the final output.

homology_arm_min = 300 #min and max base pair lengths for homology arms
homology_arm_max = 600

guide_length = 20 #length of guides in base pairs
pam_seq = 'NGG' #PAM recognition sequence "NGG" for SpCas9
cut_site_shift = 3 #Predicted nuclease cut site relative to the PAM. 3bp for SpCas9 (number of basepairs upstream of the PAM)
indel_padding = 10 #Number of basepairs padding to include between predicted cut site and nearby relevant genomic features. Default is 10
				   #This is because even if your guide is in an intron, if the cut site is too close, you can disrupt the splice site.

cut_intergenic = True #These values determine which regions can be cut and subjected to indels. By default, introns and intergenic sequences are the only ones allowed.
cut_CDS = False
cut_three_prime_UTR = False
cut_intron = True
cut_five_prime_UTR = False
cut_other_RNA = False

gibson_assembly_overhang = 20 #The number of base pair overhang used in Gibson Assembly. 20bp works just fine.

gt_space_min = 60 #The minimum and maximum base pair distances from the ends of each homology arm for genotyping primers
gt_space_max = 200 #If min is too low, you cannot get good coverage of the junciton during sequencing and the primer binding site could be disrupted by indels.
				   #If max is too high, the PCR product length may be too long to get a reliable band.

gt1_opposite = '' #Input sequences here if you have specific genotyping primers for your cargo that you want to make sure are compatible with the genomic genotyping primers.
gt2_opposite = '' #If these are blank, the cargo cloning primers will be used.
				  #For example, if your cargo is GFP and you have your own GFP primers to use with GT1 and GT2 respectively, enter them here.
				  #Otherwise, we will assume you are genotyping using GT1 + cargo_R and GT2 + cargo_F.

#EXTRACT TRANSCRIPT INFORMATION

#The chromosome number is always the third digit in a gene name.
chr_num = transcript_name[2]

print('Transcript ID: ' + transcript_name)
print('Chromosome number: ' + str(chr_num))
print('Extracting transcript features...')

#Look at the specific annotation file for this chromosome.
gff_fn = 'Arabidopsis_thaliana.TAIR10.56.chromosome.' + chr_num + '.gff3'
gff = open(gff_fn, 'r')

#Create a gene object.
gene = Gene()
for line in gff:
	#Extract information about the transcript from the annotation file.
	if transcript_name in line:
		gene.add_feature(line)
gff.close()

if gene.coding_exons:
	print('Complete.')
	print('')
else:
	print('Error! Transcript details not found in GFF file.')
	sys.exit()



#DETERMINE INSERTION COORDINATE

#Multiply by three to go from amino acid to nucleotide.
cds_insertion_location = insertion_location * 3

#If -1 was the insertion location we just want to put it at the end of the gene.
if cds_insertion_location < 0:
	print('Desired inseriton location is at the end of the coding sequence.')
	#Find the end of the coding sequence.
	#The end of this line is adding an offset and a shift of 3 to account for the stop codon.
	insertion_coordinate = gene.coding_exons[-1][gene.polarity + 1] + gene.polarity + ((-2 * gene.polarity - 1) * 3)

#If a non-negative number was given, we have to find the exact insertion site.
else:
	print('Desired insertion is after residue number: ' + str(insertion_location) + '.')
	exon_counter = 0
	#Might need a better condition now   CHECK THIS??
	while cds_insertion_location >= 0:
		if gene.coding_exons[exon_counter][2] >= cds_insertion_location:
		#If the insertion site fits in this exon, find the correct location.
			#Once again, little offsets are required to make this work for both positive and negative strand genes and make sure everything is not off by one.
			insertion_coordinate = gene.coding_exons[exon_counter][-1 * gene.polarity] + ((2 * gene.polarity + 1) * cds_insertion_location) + ((gene.polarity + 1) * -1)

			#To leave the loop.
			cds_insertion_location = -1
		else:
			#Subtract the length of this exon and move on to the next one.
			cds_insertion_location -= gene.coding_exons[exon_counter][2]
			exon_counter += 1
print('Insertion location is at coordinate ' + str(insertion_coordinate) + ' on Chromosome ' + str(chr_num) + '.')
print('')            

#To Do: Add something above to catch if the number is too high.


#PULL NUCLEOTIDE SEQUENCE

#We are only interested in the sequence flanking the insertion site so we can find guides and primers.
print('Extracting sequence flanking the insertion site...')
range_min = insertion_coordinate - local_seq_range
range_max = insertion_coordinate + local_seq_range

with open('Arabidopsis_thaliana.TAIR10.dna.chromosome.' + chr_num + '.fa') as fa:
	for record in SeqIO.parse(fa, "fasta"):
		local_seq = record.seq[range_min-1:range_max-1]

#It should already be uppercase, but we do this to be safe.
local_seq = local_seq.upper()
#Get the reverse complement of the sequence if the gene of interest is on the negative strand.
if gene.polarity == -1:
	local_seq = str(Seq(local_seq).reverse_complement())


#ENCODE FEATURES IN GENOMIC SEQUENCE

#We want an array of integers, where each integer corresponds to an individual nucleotide in the sequence.
#The value of the integer indicates what type of feature the nucleotide belongs to.

#Sequence Code:
#1: Intergenic
#2: Coding Sequence
#3: 3' UTR
#4: Intron
#5: 5' UTR
#6: Other gene type (lncRNA, miRNA, tRNA, etc.) 
print('Annotating the sequence around the insertion site...')
seq_code = np.ones(len(local_seq))
local_range = range(range_min, range_max)

#I DONT THINK THIS IS NECESSARY, DELETE IT AFTER TESTING IF IT IS NOT
# #Make sure everything is not all off by 1 or something. They are off, figure it out
# for coords in gene.five_prime_UTR:
#     five_prime_UTR_range = range(coords[0], coords[1]+1)
#     overlap = range(max(local_range[0], five_prime_UTR_range[0]), min(local_range[-1], five_prime_UTR_range[-1])+1)
#     for x in overlap:
#         seq_code[x-range_min] = 5
		
# for coords in gene.three_prime_UTR:
#     three_prime_UTR_range = range(coords[0], coords[1]+1)
#     overlap = range(max(local_range[0], three_prime_UTR_range[0]), min(local_range[-1], three_prime_UTR_range[-1])+1)
#     for x in overlap:
#         seq_code[x-range_min] = 3
		
# for coords in gene.coding_exons:
#     exon_range = range(coords[0], coords[1]+1)
#     overlap = range(max(local_range[0], exon_range[0]), min(local_range[-1], exon_range[-1])+1)
#     for x in overlap:
#         seq_code[x-range_min] = 2
		
# mRNA_range = range(gene.mRNA_start, gene.mRNA_end+1)
# overlap = range(max(local_range[0], mRNA_range[0]), min(local_range[-1], mRNA_range[-1])+1)
# for x in overlap:
#     if seq_code[x-range_min] == 1:
#         seq_code[x-range_min] = 4
	

#We want to scan through the GFF annotation file again and pull whatever information there is about the nulceotides in our local sequence
with open('Arabidopsis_thaliana.TAIR10.56.chromosome.' + chr_num + '.gff3') as gff:
	for line in gff:
		fields = line.split('\t')
		#We want to ignore the text lines at the beginning of the file
		if len(fields) > 2:
			#It is impossible to prevent all conflicts within the sequence code since there are overlapping transcripts and such, but the
			#main thing we want to avoid is issues with our specific gene. This can especially happen if there are different transcripts
			#in our gene which use sequences differentyl (i.e. retained intron). Thus we want to ignore other transcript IDs from our gene.
			if not ((transcript_name.split('.')[0] in fields[8]) and (transcript_name not in fields[8])):
				#If the feature described in the line overlaps the local sequence, we want to annotate our sequence code
				overlap = range(max(local_range[0], int(fields[3])), min(local_range[-1], int(fields[4]))+1)
				#x will be any coordinates that overlap
				for x in overlap:
					#This is so that we don't overwrite anything accidentally.
					#It is a little tricky because CDS corresponds to exons and mRNA corresponds to exons and introns,
					#so we want to make sure that introns can be overwritten but exons cannot.
					if seq_code[x-range_min] not in (2, 3, 5, 6):
						match fields[2]:
							case 'CDS':
								seq_code[x-range_min] = 2
							case 'three_prime_UTR':
								seq_code[x-range_min] = 3
							case 'five_prime_UTR':
								seq_code[x-range_min] = 5
							case 'lnc_RNA':
								seq_code[x-range_min] = 6
							case 'miRNA':
								seq_code[x-range_min] = 6
							case 'ncRNA':
								seq_code[x-range_min] = 6
							case 'snRNA':
								seq_code[x-range_min] = 6
							case 'snoRNA':
								seq_code[x-range_min] = 6
							case 'tRNA':
								seq_code[x-range_min] = 6
							case 'mRNA':
								seq_code[x-range_min] = 4
							case _:
								pass

#If the gene is on the negative strand, we flip the sequence code so that it matches up with the reverse complemented local sequence.
#From now on we will not have to use coordinates from the genome files/annotation so everything can be based off our chosen local sequence.                    
if gene.polarity == -1:
	seq_code = np.flip(seq_code)
print('Complete.')
print('')


#FIND ALL ELIGIBLE GUIDE SITES AND CHOOSE BEST GUIDE

local_center = int(len(local_seq) / 2) + (gene.polarity * 2)

#Find sites in the left arm
left_lower = local_center - homology_arm_max
left_upper = local_center - homology_arm_min

right_lower = local_center + homology_arm_min
right_upper = local_center + homology_arm_max

left_arm_seq = local_seq[left_lower:left_upper]
right_arm_seq = local_seq[right_lower:right_upper]

print('Searching for eligible guide sequences within the desired homology arm regions...')
left_guide_candidates = pull_guides(str(left_arm_seq), pam_seq, guide_length)
right_guide_candidates = pull_guides(str(right_arm_seq), pam_seq, guide_length)

guide_candidate_list = pd.DataFrame(columns = ['sequence', 'local_position', 'strand', 'arm', 'GC', 'TTTT', 'safe_to_cut'])

print('Annotating eligible guides and calculating specificity...')
#Add details about each guide manually.
#First we will add in information about the arm, the strandedness, and the start position of each guide.
for guide in left_guide_candidates:
	#We want to know the directionality of the guide. Here, + or - is relative to our local sequence, which may be flipped relative to the standardized genome.
	#We also want to know where the guide starts in our sequence so we can find its position when making primers
	if guide in str(left_arm_seq):
		guide_strand = '+'
		guide_start = str(local_seq).find(guide)
	else:
		guide_strand = '-'
		#The position for the reverse complement guides takes a little extra work.
		guide_start = (local_seq_range * 2) - str(Seq(local_seq).reverse_complement()).find(guide) - guide_length

	guide_candidate_list = pd.concat([guide_candidate_list, pd.Series({'sequence': guide, 'local_position': guide_start, 'strand': guide_strand, 'arm': 'left'}).to_frame().T], ignore_index = True)

#Do the same for the right arm guides.
for guide in right_guide_candidates:
	if guide in str(right_arm_seq):
		guide_strand = '+'
		guide_start = str(local_seq).find(guide)
	else:
		guide_strand = '-'
		guide_start = (local_seq_range * 2) - str(Seq(local_seq).reverse_complement()).find(guide) - guide_length

	guide_candidate_list = pd.concat([guide_candidate_list, pd.Series({'sequence': guide, 'local_position': guide_start, 'strand': guide_strand, 'arm': 'right'}).to_frame().T], ignore_index = True)

#Based on the types of sequence which the user has approved can be subjected to indels, make a list of allowed codes.
#This is based on the sequence code from the previous section.
allowed_codes = []
if cut_intergenic:
	allowed_codes.append(1)
if cut_CDS:
	allowed_codes.append(2)
if cut_three_prime_UTR:
	allowed_codes.append(3)
if cut_intron:
	allowed_codes.append(4)
if cut_five_prime_UTR:
	allowed_codes.append(5)
if cut_other_RNA:
	allowed_codes.append(6)

#Add more information for each guide: TTTT presence, GC content, and permission to cut.
for entry in guide_candidate_list.index:
	
	guide_seq = guide_candidate_list.loc[entry].sequence
	guide_strand = guide_candidate_list.loc[entry].strand
	guide_start = guide_candidate_list.loc[entry].local_position
	
	#Flag guide sequences with TTTT, which is a terminating signal for RNA Pol III.
	#This wouldn't be necessary if you are not transcribing gRNAs with RNA Pol III (i.e. You are not using a U6 promoter).
	if ('TTTT' in guide_seq):
		guide_candidate_list.loc[entry].TTTT = True
	else:
		guide_candidate_list.loc[entry].TTTT = False
		
	#Calculate the GC content; we assume higher GC content leads to better annealing and higher Cas9 cutting fidelity.
	guide_candidate_list.loc[entry].GC = (guide_seq.count('C') + guide_seq.count('G')) / len(guide_seq)
	
	#Determine if the cut will occur far enough away from important genomic features.
	#First find the putative cut site.
	if guide_strand == '+':
		putative_cut_site = guide_start + guide_length - cut_site_shift
	else:
		putative_cut_site = guide_start + cut_site_shift
		
	#Then get the critical window based on our padding tolerance.
	critical_window = range(putative_cut_site - indel_padding, putative_cut_site + indel_padding)

	#Compare the codes of the nucleotides in the window to those on the "allowed to cut" list.
	window_ids = set(seq_code[critical_window])
	if window_ids.issubset(allowed_codes):
		guide_candidate_list.loc[entry].safe_to_cut = True
	else:
		guide_candidate_list.loc[entry].safe_to_cut = False

#Calculate guide specificity.
#Specify which n-mers to search for. for SpCas9 and 20bp guide sites, most guide are specific to 20bp. However, it has been shown that binding to the "seed sequence"
#of roughly 8-10 bp before the PAM site can be enough to trigger nuclease activity. When this occurs at areas different from the designed guide site,
#it is known as an off-target effect. We want to choose a guide with the fewest off-targets possible.
nmer_list = [8, 10, 12, guide_length]
nmer_list.sort(reverse=True)

#Add the n-mer counts to the guide dataframe.
for num in nmer_list:
	guide_candidate_list[str(num)+'-mer_matches'] = ''

#Open the all_guides file that was generated by aphid_prep
all_guides_fn = 'all_guides_' + pam_seq + '_' + str(guide_length) + '.txt'
all_guides_file = open(all_guides_fn, 'r')
all_guides = all_guides_file.read()

#Go through each guide candidate we have so far
for entry in guide_candidate_list.index:
	#We just need the sequence of the guide
	candidate_seq = guide_candidate_list.loc[entry].sequence
	#For each n-mer we wanted, find the number of times each truncated guide matches to a PAM site in the genome.
	#The fewer matches, the better.
	for num in nmer_list:
		#Since all the guides are aligned so that the PAM is at the end of the word, we are looking for sequences that match the ends of words.
		regex_token = candidate_seq[(-1*num):] + r'\b'
		field_name = str(num)+'-mer_matches'
		#Search for the number of matches.
		matches = len(re.findall(regex_token, all_guides))
		#Add the matches to the candidate guide list.
		guide_candidate_list.loc[entry, field_name] = matches
print('All eligible guides have been annotated.')

#We want to sort the guide list so that the "best" guide is on top.
#To do this, we prioritize guide sequences without poly(T) regions and sequences that are safe to cut.
#Then we try and find the most specific guide sequence, and if there is still a tie, we choose the guide with the highest GC content.
sort_list = []
sort_order_list = []

sort_list.append('TTTT')
sort_order_list.append(True)

sort_list.append('safe_to_cut')
sort_order_list.append(False)

for num in nmer_list:
	sort_list.append(str(num) + '-mer_matches')
	sort_order_list.append(True)

sort_list.append('GC')
sort_order_list.append(False)

sorted_guide_candidate_list = guide_candidate_list.sort_values(sort_list, ascending = sort_order_list)

print('Candidate guides, sorted by quality:')
print(sorted_guide_candidate_list.to_string(index=False))
print('')

#The guide at the top of the list is the one we want to use.
best_guide = sorted_guide_candidate_list.iloc[0]

print('The best guide:')
print(best_guide.to_string())

#Provide some warnings in case the best guide is still not ideal for use.
if best_guide.TTTT == True:
	print('ALERT! Best guide is not ideal due to poly(T) presence in sequence. Consider increasing homology arm range to find more suitable guides.')
if best_guide.safe_to_cut == False:
	print('ALERT! Best guide is not ideal due to likely indels in prohibited sequence. Consider increasing homology arm range to find more suitable guides.')
if best_guide.GC < 0.3:
	print('ALERT! Best guide is not ideal due to low GC content. This may not really be an issue, but you could consider increasing homology arm range to find more suitable guides.')
field_name = str(nmer_list[1])+'-mer_matches'
if best_guide.loc[field_name] > 1:
	print('ALERT! Best guide is not ideal due to low genome specificty. To avoid off-target effects, consider increasing homology arm range to find more suitable guides.')
print('')

#DESIGN CLONING PRIMERS

#Now that we have chosen the CRISPR guide, we need to design the cloning primers to create the donor template.
#The donor template has three pieces, left arm, cargo, and right arm. Each is amplified by a separate pair of primers:
#LEFT ARM: L_F_primer_seq + L_R_primer_seq
#CARGO: cargo_F_primer_seq + cargo_R_primer_seq
#RIGHT ARM: R_F_primer_seq + R_R_primer_seq

#We will start with the simplest primers. L_R and R_F flank the insertion site, so there is little variability in what their sequence can be.
#We just extend the sequence out from the insertion site until we reach the minimum Tm.
#The minimum Tm is defined in aphid_utils but is 59C by default. Unless you are doing something very different or have issues, this does not need to change.

print('Designing cloning primers...')
#First get the bit for the L_R primer. Note: This is just the genomic portion
L_R_genomic = grow_primer_to_tm(local_seq, local_center, '-') #Remember we will RC this later

#Now get the bit for the R-F primer. Note: This is just the genomic portion.
R_F_genomic = grow_primer_to_tm(local_seq, local_center+1, '+')

#The cargo file should only contain a single sequence. Otherwise the program will use the last sequence in the file.
#Use the SeqIO docs, there is a function for just getting a single sequence or getting both.
for record in SeqIO.parse(cargo_file, "fasta"):
		cargo_seq = record.seq
		
#Get the GA overlap sequences so that the pieces of the donor template can be assembled with Gibson Assembly or overlapping PCR.
L_R_cargo = cargo_seq[0:gibson_assembly_overhang]
R_F_cargo = cargo_seq[-1*gibson_assembly_overhang:]

#Construct the final primer sequences for R_F and L_R
L_R_primer_seq = str(L_R_cargo.reverse_complement()) + str(Seq(L_R_genomic).reverse_complement())
R_F_primer_seq = str(R_F_cargo) + str(R_F_genomic)

#Now have completed the easiest primers, L-R and R-F. These are constrained by the insertion site and the cargo.

#In case the user does not have primers to amplify their cargo sequence, we can design those primers.
#First get the forward facing primer to amplify the cargo.
cargo_F_primer_seq = grow_primer_to_tm(cargo_seq, 0, '+')

#Now get the reverse primer for amplifying cargo.
cargo_R_primer_seq = grow_primer_to_tm(cargo_seq, len(cargo_seq), '-').reverse_complement()


#Next we want to make the primers for the distal ends of the homology arms, L_F and R_R.
#One of these primers will contain the gRNA sequence and PAM, and the other will just be a primer within the homology arm length range.
#Get information about our best guide so that we know which arm it is on and where it is located in the local sequence.
guide_pos = best_guide.local_position
guide_strand = best_guide.strand
guide_arm = best_guide.arm

#If the guide is on the left arm, then the guide decides the sequence of L_F.
if guide_arm == 'left': 
	#Start at the end of the guide and grow a primer out to the minimum Tm.
	if guide_strand == '+':
		primer = grow_primer_to_tm(local_seq, guide_pos, '+')
	elif guide_strand == '-':
		primer = grow_primer_to_tm(local_seq, guide_pos-len(pam_seq), '+')
	L_F_genomic = primer
	#For R_R, choose a primer that is thermodynamically ideal and is not specific to the genome, to mitigate non-specific bands.
	try:
		R_R_genomic = choose_primer(local_seq, right_lower, right_upper, '-', R_F_primer_seq).sequence
	except:
		print('Error! Could not find a genomic sequence for R_R. Try adjusting homology arm lengths or primer requirements.')
		sys.exit()

#If the guide is on the right arm, then the guide decides the sequence of R_R.
elif guide_arm == 'right':
	if guide_strand == '+':
		primer = grow_primer_to_tm(local_seq, guide_pos+len(pam_seq)+guide_length-1, '-')
	elif guide_strand == '-':
		primer = grow_primer_to_tm(local_seq, guide_pos+guide_length-1, '-')
	R_R_genomic = Seq(primer).reverse_complement()
	try:
		L_F_genomic = choose_primer(local_seq, left_lower, left_upper, '+', L_R_primer_seq).sequence
	except:
		print('Error! Could not find a genomic sequence for L_F. Try adjusting homology arm lengths or primer requirements.')
		sys.exit()
#Note: Both R_R genomic sequences have been reverse complemented at this point.

#We have the genomic components of L_F and R_R, but in order to clone them into the vector, we need the Gibson Assembly overhangs matching our plasmid.
#For pHEE-mCherry, we insert the donor in the HindIII site. The flanking sequence is always the same; I am using 20bp overhangs here.
L_plasmid_flanking = 'TAAAACGACGGCCAGTGCCA'
R_plasmid_flanking = Seq('TGTTTACACCACAATATATC').reverse_complement()

#Combind the genomic and plasmid sequence to create the full primers.
L_F_primer_seq = L_plasmid_flanking + L_F_genomic
R_R_primer_seq = R_plasmid_flanking + R_R_genomic

#We want to calculate where the cloning primers are located within the local sequence so that we can place our genotyping primers accordingly.
L_F_indices = range(str(local_seq).find(str(L_F_genomic)), str(local_seq).find(str(L_F_genomic))+len(L_F_genomic))
L_R_indices = range(str(local_seq).find(str(L_R_genomic)), str(local_seq).find(str(L_R_genomic))+len(L_R_genomic))

R_F_indices = range(str(local_seq).find(str(R_F_genomic)), str(local_seq).find(str(R_F_genomic))+len(R_F_genomic))
R_R_indices = range(str(local_seq).find(str(Seq(R_R_genomic).reverse_complement())), str(local_seq).find(str(Seq(R_R_genomic).reverse_complement()))+len(R_R_genomic))

print('Designing genotyping primers...')
#The primers for cloning the donor element are complete. Now we need two genotyping primers.
#The genotyping primers should be outside the homology arms. The distance from the homology arm ends is determined by gt_space_min and gt_space_max.
#Non-specific bands can be an issue with genotyping. These primers may not be ideal and you may need to try a few pairs of different
#genomic and cargo-specific primers to get consistent and clear results.
if gt1_opposite == '':
	gt1_opposite = str(cargo_R_primer_seq)
if gt2_opposite == '':
	gt2_opposite = str(cargo_F_primer_seq)

#Design the genotyping primers based on their approved location and chosen genotyping partner.
try:
	GT1_primer_seq = choose_primer(local_seq, min(L_F_indices)-gt_space_max, min(L_F_indices)-gt_space_min, '+', gt1_opposite).sequence
except:
	print('Error! Could not find a suitable GT1 primer sequence. Try adjusting the gt window or primer requirements.')
	sys.exit()

try:
	GT2_primer_seq = choose_primer(local_seq, max(R_R_indices)+gt_space_min, max(R_R_indices)+gt_space_max, '-', gt2_opposite).sequence
except:
	print('Error! Could not find a suitable GT2 primer sequence. Try adjusting the gt window or primer requirements.')
	sys.exit()

#Finally we need two primers for the gRNA.
#In this protocol, we anneal two primers and insert them into the plasmid via ligation.
#Using the pHEE-mCherry vector, this means we are inserting into the BsaI cut site. Therefore each primer needs specific sticky ends.
#For improved efficiency, instead of using 20bp gRNA + scaffold, we use G + 19bp gRNA + scaffold
gRNA_F_primer_seq = 'ATTG' + best_guide.sequence[1:]
gRNA_R_primer_seq = 'AAAC' + str(Seq(best_guide.sequence[1:]).reverse_complement())

print('Primer design complete.')
print('')

#We want to recalculate where the other primers are located within the local sequence so that in our final output we can highlight or underline the primer sequences in context.
GT1_indices = range(str(local_seq).find(GT1_primer_seq), str(local_seq).find(GT1_primer_seq)+len(GT1_primer_seq))
GT2_indices = range(str(local_seq).find(str(Seq(GT2_primer_seq).reverse_complement())), str(local_seq).find(str(Seq(GT2_primer_seq).reverse_complement()))+len(GT2_primer_seq))

#this is probably wrong.
gRNA_indices = range(best_guide.local_position, best_guide.local_position + guide_length)

print('Final primer sequences:')

print(project_name + '_L_F: ' + L_F_primer_seq)
print(project_name + '_L_R: ' + L_R_primer_seq)
print(project_name + '_R_F: ' + R_F_primer_seq)
print(project_name + '_R_R: ' + R_R_primer_seq)

print(cargo_name + '_F: ' + cargo_F_primer_seq)
print(cargo_name + '_R: ' + cargo_R_primer_seq)

print(project_name + '_gRNA_F: ' + gRNA_F_primer_seq)
print(project_name + '_gRNA_R: ' + gRNA_R_primer_seq)

print(project_name + '_GT1: ' + GT1_primer_seq)
print(project_name + '_GT2: ' + GT2_primer_seq)
print('')

print('Left homology arm length: ' + str(max(L_R_indices)-min(L_F_indices)+1))
print('Right homology arm length: ' + str(max(R_R_indices)-min(R_F_indices)+1))
print('Cargo sequence length: ' + str(len(cargo_seq)))
print('')


#OUTPUT FINAL RESULTS

print('Outputting final resutls...')
#Make a Word Doc with annotated sequence and other information.
design_doc = Document()
styles = design_doc.styles

paragraph = design_doc.add_paragraph()
run = paragraph.add_run('Command: python APHID.py ' + args[1] + ' ' + args[2] + ' ' + args[3] + ' ' + args[4] + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run('Transcript: ' + transcript_name + '\n' + 'Cargo: ' + cargo_name + '\n')
run.font.name = 'Courier New'

if insertion_location < 0:
	run = paragraph.add_run('Insertion location: End of coding sequence' + '\n')
else:
	run = paragraph.add_run('Insertion location: After amino acid ' + str(insertion_location) + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run('\n')


for idx,nucleotide in enumerate(str(local_seq)):
	if idx == (local_center+1):
		for nucleotide2 in str(cargo_seq):
			run = paragraph.add_run(nucleotide2)
			run.font.color.rgb = RGBColor(0, 255, 0)
			run.font.name = 'Courier New'
	run = paragraph.add_run(nucleotide)
	match seq_code[idx]:
		case 1:
			run.font.color.rgb = RGBColor(0, 0, 0)
		case 2:
			run.font.color.rgb = RGBColor(255, 165, 0)
		case 3:
			run.font.color.rgb = RGBColor(255, 0, 0)
		case 4:
			run.font.color.rgb = RGBColor(219, 123, 255)
		case 5:
			run.font.color.rgb = RGBColor(255, 0, 0)
		case 6:
			run.font.color.rgb = RGBColor(255, 165, 0)
		case _:
			pass

	if (idx in GT1_indices) or (idx in GT2_indices):
		run.font.highlight_color = 16 #Gray
	if (idx in L_F_indices) or (idx in L_R_indices):
		run.font.highlight_color = 7 #Yellow
	if (idx in R_F_indices) or (idx in R_R_indices):
		run.font.highlight_color = 3 #Turquoise
	if idx in gRNA_indices:
		run.font.underline = True


	run.font.name = 'Courier New'

run = paragraph.add_run('\n')
run = paragraph.add_run('\n')
run = paragraph.add_run('Cloning primers:\n')
run.font.name = 'Courier New'

run = paragraph.add_run(project_name + '_L_F: ' + L_F_primer_seq + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run(project_name + '_L_R: ' + L_R_primer_seq + '\n')
run.font.name = 'Courier New'

run = paragraph.add_run(project_name + '_R_F: ' + R_F_primer_seq + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run(project_name + '_R_R: ' + R_R_primer_seq + '\n')
run.font.name = 'Courier New'

run = paragraph.add_run(cargo_name + '_F: ' + cargo_F_primer_seq + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run(cargo_name + '_R: ' + cargo_R_primer_seq + '\n')
run.font.name = 'Courier New'

run = paragraph.add_run(project_name + '_gRNA_F: ' + gRNA_F_primer_seq + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run(project_name + '_gRNA_R: ' + gRNA_R_primer_seq + '\n')
run.font.name = 'Courier New'

run = paragraph.add_run('\n')
run = paragraph.add_run('Genotyping primers:\n')
run.font.name = 'Courier New'

run = paragraph.add_run(project_name + '_GT1: ' + GT1_primer_seq + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run(project_name + '_GT2: ' + GT2_primer_seq + '\n')
run.font.name = 'Courier New'

run = paragraph.add_run('\n')
run = paragraph.add_run('Left homology arm length: ' + str(max(L_R_indices)-min(L_F_indices)+1) + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run('Right homology arm length: ' + str(max(R_R_indices)-min(R_F_indices)+1) + '\n')
run.font.name = 'Courier New'
run = paragraph.add_run('Cargo sequence length: ' + str(len(cargo_seq)) + '\n')
run.font.name = 'Courier New'

paragraph.style = design_doc.styles['No Spacing']
design_doc.save(project_name + '_design.docx')

#Save the primers to a CSV file which you can open in Excel.
primer_data = [[project_name + '_L_F', L_F_primer_seq], [project_name + '_L_R', L_R_primer_seq],
				[project_name + '_R_F', R_F_primer_seq], [project_name + '_R_R', R_R_primer_seq],
				[cargo_name + '_F', cargo_F_primer_seq], [cargo_name + '_R', cargo_R_primer_seq],
				[project_name + '_gRNA_F', gRNA_F_primer_seq], [project_name + '_gRNA_R', gRNA_R_primer_seq],
				[project_name + '_GT1', GT1_primer_seq], [project_name + '_GT1', GT2_primer_seq]]
primer_df = pd.DataFrame(primer_data, columns = ['Name', 'Sequence'])
primer_df.to_csv(project_name + '_primers.csv', index=False)

#Make a FASTA file with the complete plasmid sequence.
#Note: This only works if you are using pHEE-mCherry.
whole_plasmid_seq = 'Gtttacccgccaatatatcctgtcaaacactgatagtttaaactgaaggcgggaaacgacaatctgatccaagctcaagctgctctagcattcgccattcaggctgcgcaactgttgggaagggcgatcggtgcgggcctcttcgctattacgccagctggcgaaagggggatgtgctgcaaggcgattaagttgggtaacgccagggttttcccagtcacgacgttgtaaaacgacggccagtgcca' \
					+ str(local_seq[min(L_F_indices):(max(L_R_indices)+1)]) \
					+ str(cargo_seq) \
					+ str(local_seq[min(R_F_indices):(max(R_R_indices)+1)]) \
					+ 'tgtttacaccacaatatatcctgccaAGCTGTACGGCTACATCGTGCCGAAGGATGCCCAGATCCTCGTGAACCTCTGGGCCATTGGCAGGGACCCAAACGCCTGGCAGAACGCCGATATTTTCAGCCCAGAGCGCTTCATCGGCTGCGAGATCGATGTTAAGGGCCGCGATTTCGGCCTCCTTCCATTTGGCGCTGGCCGCAGAATTTGCCCAGGCATGAATCTCGCCATCAGGATGCTCACCCTCATGCTCGCCACACTCCTCCAGTTCTTCAACTGGAAGCTCGAAGGCGACATCTCCCCGAAGGACCTCGACATGGACGgtttacccgccaatatatcctgtcatcgacttgccttccgcacaatacatcatttcttcttagctttttttcttcttcttcgttcatacagtttttttttgtttatcagcttacattttcttgaaccgtagctttcgttttcttctttttaactttccattcggagtttttgtatcttgtttcatagtttgtcccaggattagaatgattaggcatcgaaccttcaagaatttgattgaataaaacatcttcattcttaagatatgaagataatcttcaaaaggcccctgggaatctgaaagaagagaagcaggcccatttatatgggaaagaacaatagtatttcttatataggcccatttaagttgaaaacaatcttcaaaagtcccacatcgcttagataagaaaacgaagctgagtttatatacagctagagtcgaagtagtgattg' \
					+ best_guide.sequence \
					+ 'gttttagagctagaaatagcaagttaaaataaggctagtccgttatcaacttgaaaaagtggcaccgagtcggtgcttttttttgcaaaattttccagatcgatttcttcttcctctgttcttcggcgttcaatttctggggttttctcttcgttttctgtaactgaaacctaaaatttgacctaaaaaaaatctcaaataatatgattcagtggttttgtacttttcagttagttgagttttgcagttccgatgagataaaccaataccatggttatactagtgaataaaagcatttgcgtttggtttatcattgcgtttatacaaggacagagatccactgagctggaatagcttaaaaccattatcagaacaaaataaaccattttttgttaagaatcagagcatagtaaacaacagaaacaacctaagagaggtaacttgtccaagaagatagctaattatatctattttataaaagttatcatagtttgtaagtcacaaaagatgcaaataacagagaaactaggagacttgagaatatacattcttgtatatttgtattcgagattgtgaaaatttgaccataagtttaaattcttaaaaagatatatctgatctaggtgatggttatagactgtaattttaccacatgtttaatgatggatagtgacacacatgacacatcgacaacactatagcatcttatttagattacaacatgaaatttttctgtaatacatgtctttgtacataatttaaaagtaattcctaagaaatatatttatacaaggagtttaaagaaaacatagcataaagttcaatgagtagtaaaaaccatatacagtatatagcataaagttcaatgagtttattacaaaagcattggttcactttctgtaacacgacgttaaaccttcgtctccaataggagcgctactgattcaacatgccaatatatactaaatacgtttctacagtcaaatgctttaacgtttcatg' \
					+ 'attaagtgactatttaccgtcaatcctttcccattcctcccactaatccaactttttaattactcttaaatcaccactaagctagtaacgcctatcatgaattagctctactaaatctagcaacctttcaaatttgcagtattgcaggtgtctctgtgtctttaaaatagttgccttatgatttcttcggtttcaagatgatcaaatagttatagatttcatgctcacacatgctcattagatgtgtacatactttacttacccaaatctattttctcgcaaagattttgatggtaaagctgatttggttctattgaactaaatcaaacgagtttcagactgagtgattctaatccggcccattagcccctaaacagacccactaattacgcagcttttaatagagtaattacacctagtttacccactaaaccactaagcactaattatctcacaatctaatgagcttccctcgtaattacttgggctttcactctaccatttatttgtaacagtcaagtctctactgtctctatataaactctctaaagttaacacacaattctcatcacaaacaaatcaaccaaagcaacttctactctttcttctttcgaccttatcaatctgttgagaaatctagatggattacaaggaccacgacggggattacaaggaccacgacattgattacaaggatgatgatgacaagatggctccgaagaagaagaggaaggttggcatccacggggtgccagctgctgacaagaagtactcgatcggcctcgatattgggactaactctgttggctgggccgtgatcaccgacgagtacaaggtgccctcaaagaagttcaaggtcctgggcaacaccgatcggcattccatcaagaagaatctcattggcgctctcctgttcgacagcggcgagacggctgaggctacgcggctcaagcgcaccgcccgcaggcggtacacgcgcaggaagaatcgcatctgctacctgcaggagattttctccaacgag' \
					+ 'atggcgaaggttgacgattctttcttccacaggctggaggagtcattcctcgtggaggaggataagaagcacgagcggcatccaatcttcggcaacattgtcgacgaggttgcctaccacgagaagtaccctacgatctaccatctgcggaagaagctcgtggactccacagataaggcggacctccgcctgatctacctcgctctggcccacatgattaagttcaggggccatttcctgatcgagggggatctcaacccggacaatagcgatgttgacaagctgttcatccagctcgtgcagacgtacaaccagctcttcgaggagaaccccattaatgcgtcaggcgtcgacgcgaaggctatcctgtccgctaggctctcgaagtctcggcgcctcgagaacctgatcgcccagctgccgggcgagaagaagaacggcctgttcgggaatctcattgcgctcagcctggggctcacgcccaacttcaagtcgaatttcgatctcgctgaggacgccaagctgcagctctccaaggacacatacgacgatgacctggataacctcctggcccagatcggcgatcagtacgcggacctgttcctcgctgccaagaatctgtcggacgccatcctcctgtctgatattctcagggtgaacaccgagattacgaaggctccgctctcagcctccatgatcaagcgctacgacgagcaccatcaggatctgaccctcctgaaggcgctggtcaggcagcagctccccgagaagtacaaggagatcttcttcgatcagtcgaagaacggctacgctgggtacattgacggcggggcctctcaggaggagttctacaagttcatcaagccgattctggagaagatggacggcacggaggagctgctggtgaagctcaatcgcgaggacctcctgaggaagcagcggacattcgataacggcagcatcccacaccagattcatctcggggagctgcacgctatcctgaggaggcaggaggacttctacc' \
					+ 'ctttcctcaaggataaccgcgagaagatcgagaagattctgactttcaggatcccgtactacgtcggcccactcgctaggggcaactcccgcttcgcttggatgacccgcaagtcagaggagacgatcacgccgtggaacttcgaggaggtggtcgacaagggcgctagcgctcagtcgttcatcgagaggatgacgaatttcgacaagaacctgccaaatgagaaggtgctccctaagcactcgctcctgtacgagtacttcacagtctacaacgagctgactaaggtgaagtatgtgaccgagggcatgaggaagccggctttcctgtctggggagcagaagaaggccatcgtggacctcctgttcaagaccaaccggaaggtcacggttaagcagctcaaggaggactacttcaagaagattgagtgcttcgattcggtcgagatctctggcgttgaggaccgcttcaacgcctccctggggacctaccacgatctcctgaagatcattaaggataaggacttcctggacaacgaggagaatgaggatatcctcgaggacattgtgctgacactcactctgttcgaggaccgggagatgatcgaggagcgcctgaagacttacgcccatctcttcgatgacaaggtcatgaagcagctcaagaggaggaggtacaccggctgggggaggctgagcaggaagctcatcaacggcattcgggacaagcagtccgggaagacgatcctcgacttcctgaagagcgatggcttcgcgaaccgcaatttcatgcagctgattcacgatgacagcctcacattcaaggaggatatccagaaggctcaggtgagcggccagggggactcgctgcacgagcatatcgcgaacctcgctggctcgccagctatcaagaaggggattctgcagaccgtgaaggttgtggacgagctggtgaaggtcatgggcaggcacaagcctgagaacatcgtcattgagatggcccgggagaatcagaccacgcagaa' \
					+ 'gggccagaagaactcacgcgagaggatgaagaggatcgaggagggcattaaggagctggggtcccagatcctcaaggagcacccggtggagaacacgcagctgcagaatgagaagctctacctgtactacctccagaatggccgcgatatgtatgtggaccaggagctggatattaacaggctcagcgattacgacgtcgatcatatcgttccacagtcattcctgaaggatgactccattgacaacaaggtcctcaccaggtcggacaagaaccggggcaagtctgataatgttccttcagaggaggtcgttaagaagatgaagaactactggcgccagctcctgaatgccaagctgatcacgcagcggaagttcgataacctcacaaaggctgagaggggcgggctctctgagctggacaaggcgggcttcatcaagaggcagctggtcgagacacggcagatcactaagcacgttgcgcagattctcgactcacggatgaacactaagtacgatgagaatgacaagctgatccgcgaggtgaaggtcatcaccctgaagtcaaagctcgtctccgacttcaggaaggatttccagttctacaaggttcgggagatcaacaattaccaccatgcccatgacgcgtacctgaacgcggtggtcggcacagctctgatcaagaagtacccaaagctcgagagcgagttcgtgtacggggactacaaggtttacgatgtgaggaagatgatcgccaagtcggagcaggagattggcaaggctaccgccaagtacttcttctactctaacattatgaatttcttcaagacagagatcactctggccaatggcgagatccggaagcgccccctcatcgagacgaacggcgagacgggggagatcgtgtgggacaagggcagggatttcgcgaccgtcaggaaggttctctccatgccacaagtgaatatcgtcaagaagacagaggtccagactggcgggttctctaaggagtcaattctgcctaag' \
					+ 'cggaacagcgacaagctcatcgcccgcaagaaggactgggatccgaagaagtacggcgggttcgacagccccactgtggcctactcggtcctggttgtggcgaaggttgagaagggcaagtccaagaagctcaagagcgtgaaggagctgctggggatcacgattatggagcgctccagcttcgagaagaacccgatcgatttcctggaggcgaagggctacaaggaggtgaagaaggacctgatcattaagctccccaagtactcactcttcgagctggagaacggcaggaagcggatgctggcttccgctggcgagctgcagaaggggaacgagctggctctgccgtccaagtatgtgaacttcctctacctggcctcccactacgagaagctcaagggcagccccgaggacaacgagcagaagcagctgttcgtcgagcagcacaagcattacctcgacgagatcattgagcagatttccgagttctccaagcgcgtgatcctggccgacgcgaatctggataaggtcctctccgcgtacaacaagcaccgcgacaagccaatcagggagcaggctgagaatatcattcatctcttcaccctgacgaacctcggcgcccctgctgctttcaagtacttcgacacaactatcgatcgcaagaggtacacaagcactaaggaggtcctggacgcgaccctcatccaccagtcgattaccggcctctacgagacgcgcatcgacctgtctcagctcgggggcgacaagcggccagcggcgacgaagaaggcggggcaggcgaagaagaagaagtgagctcagagctttcgttcgtatcatcggtttcgacaacgttcgtcaagttcaatgcatcagtttcattgcgcacacaccagaatcctactgagtttgagtattatggcattgggaaaactgtttttcttgtaccatttgttgtgcttgtaatttactgtgttttttattcggttttcgctatcgaactgtgaaatggaaatggatggag' \
					+ 'aagagttaatgaatgatatggtccttttgttcattctcaaattaatattatttgttttttctcttatttgttgtgtgttgaatttgaaattataagagatatgcaaacattttgttttgagtaaaaatgtgtcaaatcgtggcctctaatgaccgaagttaatatgaggagtaaaacacttgtagttgtaccattatgcttattcactaggcaacaaatatattttcagacctagaaaagctgcaaatgttactgaatacaagtatgtcctcttgtgttttagacatttatgaactttcctttatgtaattttccagaatccttgtcagattctaatcattgctttataattatagttatactcatggatttgtagttgagtatgaaaatattttttaatgcattttatgacttgctaagctggcacaactatatttccaacatcactagctaccatcaaaagattgacttctcatcttactcgattgaaaccaaattaacatagggtttttatttaaataaaagtttaaccttctttttaaaaaattgttcatagtgtcatgtcagaacaagagctacaaatcacacatagcatgcataagcggagctatgatgagtggtattgttttgttcgtcacttgtcactcttttccaacacataatcccgacaacaacgtaagagcatctctctctctccacacacactcatgcatgcatgcattcttacacgtgattgccatgcaaatctcctttctcacctataaatacaaaccaacccttcactacactcttcactcaaaccaaaacaagaaaacatacacaaatagcaaaacggtaccAACAATGGATAACATGGCCATCATCAAGGAGTTCATGCGCTTCAAGGTGCACATGGAGGGCTCCGTGAACGGCCACGAGTTCGAGATCGAGGGCGAGGGCGAGGGCCGCCCCTACGAGGGCACCCAGACCGCCAAGCTGAAGGTGACCAAGGGTGGCCCCCTGCCCTTCGCCTGGGACATCCTGT' \
					+ 'CCCCTCAGTTCATGTACGGCTCCAAGGCCTACGTGAAGCACCCCGCCGACATCCCCGACTACTTGAAGCTGTCCTTCCCCGAGGGCTTCAAGTGGGAGCGCGTGATGAACTTCGAGGACGGCGGCGTGGTGACCGTGACCCAGGACTCCTCCCTGCAGGACGGCGAGTTCATCTACAAGGTGAAGCTGCGCGGCACCAACTTCCCCTCCGACGGCCCCGTAATGCAGAAGAAGACCATGGGCTGGGAGGCCTCCTCCGAGCGGATGTACCCCGAGGACGGCGCCCTGAAGGGCGAGATCAAGCAGAGGCTGAAGCTGAAGGACGGCGGCCACTACGACGCTGAGGTCAAGACCACCTACAAGGCCAAGAAGCCCGTGCAGCTGCCCGGCGCCTACAACGTCAACATCAAGTTGGACATCACCTCCCACAACGAGGACTACACCATCGTGGAACAGTACGAACGCGCCGAGGGCCGCCACTCCACCGGCGGCATGGACGAGCTGTACAAGTGAgagctcttggactcccatgttggcaaaggcaaccaaacaaacaatgaatgatccgctcctgcatatggggcggtttgagtatttcaactgccatttgggctgaattgaagacatgctcctgtcagaaattccgtgatcttactcaatattcagtaatctcggccaatatcctaaatgtgcgtggctttatctgtctttgtattgtttcatcaattcatgtaacgtttgcttttcttatgaattttcaaataaattatcGtattaattgattgacaacgaattcgtaatcatgtcatagctgtttcctgtgtgaaattgttatccgctcacaattccacacaacatacgagccggaagcataaagtgtaaagcctggggtgcctaatgagtgagctaactcacattaattgcgttgcgctcactgcccgctttccagtcgggaaacctgtcgtgccagctgcattaatgaatcggccaacgcgcggggagaggcggtttgcgtattggctaga' \
					+ 'gcagcttgccaacatggtggagcacgacactctcgtctactccaagaatatcaaagatacagtctcagaagaccaaagggctattgagacttttcaacaaagggtaatatcgggaaacctcctcggattccattgcccagctatctgtcacttcatcaaaaggacagtagaaaaggaaggtggcacctacaaatgccatcattgcgataaaggaaaggctatcgttcaagatgcctctgccgacagtggtcccaaagatggacccccacccacgaggagcatcgtggaaaaagaagacgttccaaccacgtcttcaaagcaagtggattgatgtgataacatggtggagcacgacactctcgtctactccaagaatatcaaagatacagtctcagaagaccaaagggctattgagacttttcaacaaagggtaatatcgggaaacctcctcggattccattgcccagctatctgtcacttcatcaaaaggacagtagaaaaggaaggtggcacctacaaatgccatcattgcgataaaggaaaggctatcgttcaagatgcctctgccgacagtggtcccaaagatggacccccacccacgaggagcatcgtggaaaaagaagacgttccaaccacgtcttcaaagcaagtggattgatgtgatatctccactgacgtaagggatgacgcacaatcccactatccttcgcaagaccttcctctatataaggaagttcatttcatttggagaggacacgctgaaatcaccagtctctctctacaaatctatctctctcgagctttcgcagatcccggggggcaatgagatatgaaaaagcctgaactcaccgcgacgtctgtcgagaagtttctgatcgaaaagttcgacagcgtctccgacctgatgcagctctcggagggcgaagaatctcgtgctttcagcttcgatgtaggagggcgtggatatgtcctgcgggtaaatagctgcgccgatggtttctacaaagatcgttatgtttatcgg' \
					+ 'cactttgcatcggccgcgctcccgattccggaagtgcttgacattggggagtttagcgagagcctgacctattgcatctcccgccgtgcacagggtgtcacgttgcaagacctgcctgaaaccgaactgcccgctgttctacaaccggtcgcggaggctatggatgcgatcgctgcggccgatcttagccagacgagcgggttcggcccattcggaccgcaaggaatcggtcaatacactacatggcgtgatttcatatgcgcgattgctgatccccatgtgtatcactggcaaactgtgatggacgacaccgtcagtgcgtccgtcgcgcaggctctcgatgagctgatgctttgggccgaggactgccccgaagtccggcacctcgtgcacgcggatttcggctccaacaatgtcctgacggacaatggccgcataacagcggtcattgactggagcgaggcgatgttcggggattcccaatacgaggtcgccaacatcttcttctggaggccgtggttggcttgtatggagcagcagacgcgctacttcgagcggaggcatccggagcttgcaggatcgccacgactccgggcgtatatgctccgcattggtcttgaccaactctatcagagcttggttgacggcaatttcgatgatgcagcttgggcgcagggtcgatgcgacgcaatcgtccgatccggagccgggactgtcgggcgtacacaaatcgcccgcagaagcgcggccgtctggaccgatggctgtgtagaagtactcgccgatagtggaaaccgacgccccagcactcgtccgagggcaaagaaatagagtagatgccgaccggatctgtcgatcgacaagctcgagtttctccataataatgtgtgagtagttcccagataagggaattagggttcctatagggtttcgctcatgtgttgagcatataagaaacccttagtatgtatttgtatttgtaaaatacttctatcaataaaatttctaattcctaaaaccaaaa' \
					+ 'tccagtactaaaatccagatcccccgaattaattcggcgttaattcagtacattaaaaacgtccgcaatgtgttattaagttgtctaagcgtcaatttgtttacaccacaatatatcctgccaccagccagccaacagctccccgaccggcagctcggcacaaaatcaccactcgatacaggcagcccatcagtccgggacggcgtcagcgggagagccgttgtaaggcggcagactttgctcatgttaccgatgctattcggaagaacggcaactaagctgccgggtttgaaacacggatgatctcgcggagggtagcatgttgattgtaacgatgacagagcgttgctgcctgtgatcaccgcggtttcaaaatcggctccgtcgatactatgttatacgccaactttgaaaacaactttgaaaaagctgttttctggtatttaaggttttagaatgcaaggaacagtgaattggagttcgtcttgttataattagcttcttggggtatctttaaatactgtagaaaagaggaaggaaataataaatggctaaaatgagaatatcaccggaattgaaaaaactgatcgaaaaataccgctgcgtaaaagatacggaaggaatgtctcctgctaaggtatataagctggtgggagaaaatgaaaacctatatttaaaaatgacggacagccggtataaagggaccacctatgatgtggaacgggaaaaggacatgatgctatggctggaaggaaagctgcctgttccaaaggtcctgcactttgaacggcatgatggctggagcaatctgctcatgagtgaggccgatggcgtcctttgctcggaagagtatgaagatgaacaaagccctgaaaagattatcgagctgtatgcggagtgcatcaggctctttcactccatcgacatatcggattgtccctatacgaatagcttagacagccgcttagccgaattggattacttactgaataacgatctggccgatgtggattgcgaaaactgg' \
					+ 'gaagaagacactccatttaaagatccgcgcgagctgtatgattttttaaagacggaaaagcccgaagaggaacttgtcttttcccacggcgacctgggagacagcaacatctttgtgaaagatggcaaagtaagtggctttattgatcttgggagaagcggcagggcggacaagtggtatgacattgccttctgcgtccggtcgatcagggaggatatcggggaagaacagtatgtcgagctattttttgacttactggggatcaagcctgattgggagaaaataaaatattatattttactggatgaattgttttagtacctagaatgcatgaccaaaatcccttaacgtgagttttcgttccactgagcgtcagaccccgtagaaaagatcaaaggatcttcttgagatcctttttttctgcgcgtaatctgctgcttgcaaacaaaaaaaccaccgctaccagcggtggtttgtttgccggatcaagagctaccaactctttttccgaaggtaactggcttcagcagagcgcagataccaaatactgtccttctagtgtagccgtagttaggccaccacttcaagaactctgtagcaccgcctacatacctcgctctgctaatcctgttaccagtggctgctgccagtggcgataagtcgtgtcttaccgggttggactcaagacgatagttaccggataaggcgcagcggtcgggctgaacggggggttcgtgcacacagcccagcttggagcgaacgacctacaccgaactgagatacctacagcgtgagctatgagaaagcgccacgcttcccgaagggagaaaggcggacaggtatccggtaagcggcagggtcggaacaggagagcgcacgagggagcttccagggggaaacgcctggtatctttatagtcctgtcgggtttcgccacctctgacttgagcgtcgatttttgtgatgctcgtcaggggggcggagcctatggaaaaacgccagcaacgcggcctttttacggttcc' \
					+ 'tggccttttgctggccttttgctcacatgttctttcctgcgttatcccctgattctgtggataaccgtattaccgcctttgagtgagctgataccgctcgccgcagccgaacgaccgagcgcagcgagtcagtgagcgaggaagcggaagagcgcctgatgcggtattttctccttacgcatctgtgcggtatttcacaccgcatatggtgcactctcagtacaatctgctctgatgccgcatagttaagccagtatacactccgctatcgctacgtgactgggtcatggctgcgccccgacacccgccaacacccgctgacgcgccctgacgggcttgtctgctcccggcatccgcttacagacaagctgtgaccgtctccgggagctgcatgtgtcagaggttttcaccgtcatcaccgaaacgcgcgaggcagggtgccttgatgtgggcgccggcggtcgagtggcgacggcgcggcttgtccgcgccctggtagattgcctggccgtaggccagccatttttgagcggccagcggccgcgataggccgacgcgaagcggcggggcgtagggagcgcagcgaccgaagggtaggcgctttttgcagctcttcggctgtgcgctggccagacagttatgcacaggccaggcgggttttaagagttttaataagttttaaagagttttaggcggaaaaatcgccttttttctcttttatatcagtcacttacatgtgtgaccggttcccaatgtacggctttgggttcccaatgtacgggttccggttcccaatgtacggctttgggttcccaatgtacgtgctatccacaggaaacagaccttttcgacctttttcccctgctagggcaatttgccctagcatctgctccgtacattaggaaccggcggatgcttcgccctcgatcaggttgcggtagcgcatgactaggatcgggccagcctgccccgcctcctccttcaaatcgtactccggcaggtcatttgacccgatcagcttgcg' \
					+ 'cacggtgaaacagaacttcttgaactctccggcgctgccactgcgttcgtagatcgtcttgaacaaccatctggcttctgccttgcctgcggcgcggcgtgccaggcggtagagaaaacggccgatgccgggatcgatcaaaaagtaatcggggtgaaccgtcagcacgtccgggttcttgccttctgtgatctcgcggtacatccaatcagctagctcgatctcgatgtactccggccgcccggtttcgctctttacgatcttgtagcggctaatcaaggcttcaccctcggataccgtcaccaggcggccgttcttggccttcttcgtacgctgcatggcaacgtgcgtggtgtttaaccgaatgcaggtttctaccaggtcgtctttctgctttccgccatcggctcgccggcagaacttgagtacgtccgcaacgtgtggacggaacacgcggccgggcttgtctcccttcccttcccggtatcggttcatggattcggttagatgggaaaccgccatcagtaccaggtcgtaatcccacacactggccatgccggccggccctgcggaaacctctacgtgcccgtctggaagctcgtagcggatcacctcgccagctcgtcggtcacgcttcgacagacggaaaacggccacgtccatgatgctgcgactatcgcgggtgcccacgtcatagagcatcggaacgaaaaaatctggttgctcgtcgcccttgggcggcttcctaatcgacggcgcaccggctgccggcggttgccgggattctttgcggattcgatcagcggccgcttgccacgattcaccggggcgtgcttctgcctcgatgcgttgccgctgggcggcctgcgcggccttcaacttctccaccaggtcatcacccagcgccgcgccgatttgtaccgggccggatggtttgcgaccgctcacgccgattcctcgggcttgggggttccagtgccattgcagggccggcagacaacccagccgcttacgcctggccaac' \
					+ 'cgcccgttcctccacacatggggcattccacggcgtcggtgcctggttgttcttgattttccatgccgcctcctttagccgctaaaattcatctactcatttattcatttgctcatttactctggtagctgcgcgatgtattcagatagcagctcggtaatggtcttgccttggcgtaccgcgtacatcttcagcttggtgtgatcctccgccggcaactgaaagttgacccgcttcatggctggcgtgtctgccaggctggccaacgttgcagccttgctgctgcgtgcgctcggacggccggcacttagcgtgtttgtgcttttgctcattttctctttacctcattaactcaaatgagttttgatttaatttcagcggccagcgcctggacctcgcgggcagcgtcgccctcgggttctgattcaagaacggttgtgccggcggcggcagtgcctgggtagctcacgcgctgcgtgatacgggactcaagaatgggcagctcgtacccggccagcgcctcggcaacctcaccgccgatgcgcgtgcctttgatcgcccgcgacacgacaaaggccgcttgtagccttccatccgtgacctcaatgcgctgcttaaccagctccaccaggtcggcggtggcccatatgtcgtaagggcttggctgcaccggaatcagcacgaagtcggctgccttgatcgcggacacagccaagtccgccgcctggggcgctccgtcgatcactacgaagtcgcgccggccgatggccttcacgtcgcggtcaatcgtcgggcggtcgatgccgacaacggttagcggttgatcttcccgcacggccgcccaatcgcgggcactgccctggggatcggaatcgactaacagaacatcggccccggcgagttgcagggcgcgggctagatgggttgcgatggtcgtcttgcctgacccgcctttctggttaagtacagcgataaccttcatgcgttccccttgcgtatttgtttatttactcatcgcatcata' \
					+ 'tacgcagcgaccgcatgacgcaagctgttttactcaaatacacatcacctttttagacggcggcgctcggtttcttcagcggccaagctggccggccaggccgccagcttggcatcagacaaaccggccaggatttcatgcagccgcacggttgagacgtgcgcgggcggctcgaacacgtacccggccgcgatcatctccgcctcgatctcttcggtaatgaaaaacggttcgtcctggccgtcctggtgcggtttcatgcttgttcctcttggcgttcattctcggcggccgccagggcgtcggcctcggtcaatgcgtcctcacggaaggcaccgcgccgcctggcctcggtgggcgtcacttcctcgctgcgctcaagtgcgcggtacagggtcgagcgatgcacgccaagcagtgcagccgcctctttcacggtgcggccttcctggtcgatcagctcgcgggcgtgcgcgatctgtgccggggtgagggtagggcgggggccaaacttcacgcctcgggccttggcggcctcgcgcccgctccgggtgcggtcgatgattagggaacgctcgaactcggcaatgccggcgaacacggtcaacaccatgcggccggccggcgtggtggtgtcggcccacggctctgccaggctacgcaggcccgcgccggcctcctggatgcgctcggcaatgtccagtaggtcgcgggtgctgcgggccaggcggtctagcctggtcactgtcacaacgtcgccagggcgtaggtggtcaagcatcctggccagctccgggcggtcgcgcctggtgccggtgatcttctcggaaaacagcttggtgcagccggccgcgtgcagttcggcccgttggttggtcaagtcctggtcgtcggtgctgacgcgggcatagcccagcaggccagcggcggcgctcttgttcatggcgtaatgtctccggttctagtcgcaagtattctactttatgcgactaaaacacgcgacaagaaaacgccaggaaaaggg' \
					+ 'cagggcggcagcctgtcgcgtaacttaggacttgtgcgacatgtcgttttcagaagacggctgcactgaacgtcagaagccgactgcactatagcagcggaggggttggatcaaagtactttgatcccgaggggaaccctgtggttggcatgcacatacaaatggacgaacggataaaccttttcacgcccttttaaatatccgttattctaataaacgctcttttctcttag'
with open(project_name + '_plasmid.fa', 'w') as f:
	f.write('>pHEE_mCherry_' + project_name + '\n')
	f.write(whole_plasmid_seq)

print('Design complete.')
